import os
import tempfile
from pdf2docx import Converter
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml.html import builder as E
from lxml import html
from .utils import get_logger
from .image_handler import extract_images_from_doc, get_image_position
from .styles import HTML_HEAD_TEMPLATE, HTML_TAIL, DEFAULT_FONT_FAMILY
from .processor import process_text_node
import multiprocessing

def run_has_line_break(r):
    try:
        ns_w = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        brs = r._element.findall('.//w:br', ns_w)
        return len(brs) > 0
    except Exception:
        return False

def get_dominant_font_size(doc):
    """
    Scans the document to find the most common font size.
    Returns a string like '12pt'.
    """
    from collections import Counter
    sizes = Counter()
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.size:
                sizes[run.font.size.pt] += 1
                
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                         if run.font.size:
                            sizes[run.font.size.pt] += 1
                            
    if not sizes:
        return "11pt" # Default if no size found
        
    # Get most common
    most_common_pt = sizes.most_common(1)[0][0]
    # Round to nearest integer or 0.5 to clean it up? 
    # Let's keep it as is but maybe ensure it's a string
    return f"{most_common_pt}pt"

def build_paragraph_element(paragraph, images_dict, bold_ratio=0.5):
    # Determine alignment
    style_align = ""
    if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        style_align = "text-align: center;"
    elif paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        style_align = "text-align: right;"
    
    # Outer div
    p_div = E.DIV(CLASS("paragraph"))
    if style_align:
        p_div.set("style", style_align)
        
    for run in paragraph.runs:
        # Check for images in run
        ns_w = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        drawing_elements = run._element.findall('.//w:drawing', ns_w)
        
        # Handle images
        for drawing in drawing_elements:
            ns_a = {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}
            blip = drawing.find('.//a:blip', ns_a)
            if blip is not None:
                ns_rel = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
                embed_attr = f'{{{ns_rel}}}embed'
                image_rel_id = blip.get(embed_attr)
                
                if image_rel_id and image_rel_id in images_dict:
                    img_src = images_dict[image_rel_id]
                    position = get_image_position(drawing)
                    
                    img_tag = E.IMG(src=img_src, alt="Document Image")
                    container = E.DIV(CLASS("image-container"))
                    container.append(img_tag)
                    
                    if position:
                        # User reported overlaps with absolute positioning. 
                        # We switch to relative/block layout to ensure no text overlap.
                        # style_pos = f"position: absolute; left: {position['x']}in; top: {position['y']}in;"
                        # container.set("style", style_pos)
                        pass
                        
                    p_div.append(container)
        
        # Handle Text
        text = run.text
        if text:
            # Build style string
            styles = []
            if run.bold: styles.append('font-weight: bold')
            if run.italic: styles.append('font-style: italic')
            if run.underline: styles.append('text-decoration: underline')
            if run.font.size: styles.append(f'font-size: {run.font.size.pt}pt')
            
            # span = E.SPAN(text)
            # if styles:
            #     span.set("style", "; ".join(styles))
            # p_div.append(span)
            
            # Apply bionic reading here!
            bionic_html = process_text_node(text, bold_ratio)
            
            # If bionic_html is None/Empty, handle it
            if not bionic_html:
                 span = E.SPAN(text)
                 if styles: span.set("style", "; ".join(styles))
                 p_div.append(span)
            else:
                 # It returns a string with spans. We need to parse it or inject it.
                 # Since it contains multiple elements (spans), we can parse it as a fragment.
                 try:
                     # Wrap in a span to hold styles
                     container_span = E.SPAN()
                     if styles:
                         container_span.set("style", "; ".join(styles))
                     
                     # We can't easily parse fragments with lxml.builder.
                     # But we can parse with lxml.html.fragment_fromstring
                     # However, fragment_fromstring might return a single element or fail on multiple.
                     # Let's wrap bionic_html in a dummy tag to parse multiple children
                     dummy = html.fromstring(f"<div>{bionic_html}</div>")
                     for child in dummy:
                         container_span.append(child)
                         # Also handle text tails? 
                         # lxml handles tails automatically if we append elements.
                         # But fromstring might have text nodes mixed with elements.
                         # Actually `process_text_node` returns string with spans and text.
                         # This needs careful handling.
                         
                     # Easier way: Just set innerHTML? lxml doesn't support that easily.
                     # Let's trust `process_text_node` returns valid HTML fragment.
                     # We can iterate over the parsed dummy div's content.
                     if dummy.text:
                         container_span.text = dummy.text
                     for child in dummy:
                         container_span.append(child)
                         
                     p_div.append(container_span)
                 except Exception as e:
                     # Fallback
                     span = E.SPAN(text)
                     if styles: span.set("style", "; ".join(styles))
                     p_div.append(span)

        if run_has_line_break(run):
             p_div.append(E.BR())

    return p_div

def CLASS(cls):
    return {"class": cls}

def process_table(table, images_dict, bold_ratio=0.5):
    tbl = E.TABLE()
    for row in table.rows:
        tr = E.TR()
        for cell in row.cells:
            td = E.TD()
            for paragraph in cell.paragraphs:
                p_elem = build_paragraph_element(paragraph, images_dict, bold_ratio)
                td.append(p_elem)
            tr.append(td)
        tbl.append(tr)
    return tbl

def pdf_to_html(pdf_path, font_family=DEFAULT_FONT_FAMILY, bold_ratio=0.5):
    logger = get_logger()
    
    # Create temp dir for docx
    with tempfile.TemporaryDirectory() as temp_dir:
        docx_path = os.path.join(temp_dir, "temp.docx")
        logger.info(f"Converting PDF to DOCX: {docx_path}")
        
        try:
            cv = Converter(pdf_path)
            # Enable multiprocessing
            cpu_count = multiprocessing.cpu_count()
            cv.convert(docx_path, multi_processing=True, cpu_count=cpu_count)
            cv.close()
        except Exception as e:
            logger.error(f"pdf2docx failed: {e}")
            raise

        logger.info("Parsing DOCX...")
        doc = Document(docx_path)
        images_dict = extract_images_from_doc(doc)
        
        # Detect font size
        font_size = get_dominant_font_size(doc)
        logger.info(f"Detected dominant font size: {font_size}")

        # Start building HTML
        # We'll build the body content as a list of elements
        body_content = []
        
        # doc.paragraphs only gives top-level paragraphs. 
        # Tables are separate. We need to iterate over body elements to keep order.
        
        # Map paragraphs to index for easy lookup
        paragraph_map = {p._element: p for p in doc.paragraphs}
        table_map = {t._element: t for t in doc.tables}
        
        for element in doc.element.body:
             if element.tag.endswith('p'):
                 if element in paragraph_map:
                     p_elem = build_paragraph_element(paragraph_map[element], images_dict, bold_ratio)
                     body_content.append(p_elem)
             elif element.tag.endswith('tbl'):
                 if element in table_map:
                     t_elem = process_table(table_map[element], images_dict, bold_ratio)
                     body_content.append(t_elem)

        # Chunk the content
        CHUNK_SIZE = 100 # Elements per chunk
        chunks = []
        
        # Prepare template parts
        head = HTML_HEAD_TEMPLATE.format(
            font_family=font_family, 
            font_size=font_size,
            bold_weight="700", 
            normal_weight="400"
        )
        tail = HTML_TAIL
        
        current_chunk_elements = []
        for i, elem in enumerate(body_content):
            current_chunk_elements.append(elem)
            if len(current_chunk_elements) >= CHUNK_SIZE:
                 # Render this chunk
                 chunk_html = head
                 for e in current_chunk_elements:
                     chunk_html += html.tostring(e, encoding='unicode')
                 chunk_html += tail
                 chunks.append(chunk_html)
                 current_chunk_elements = []
                 
        # Remaining elements
        if current_chunk_elements:
             chunk_html = head
             for e in current_chunk_elements:
                 chunk_html += html.tostring(e, encoding='unicode')
             chunk_html += tail
             chunks.append(chunk_html)
             
        # Guard against empty content
        if not chunks:
            chunks.append(head + tail)
            
        return chunks
